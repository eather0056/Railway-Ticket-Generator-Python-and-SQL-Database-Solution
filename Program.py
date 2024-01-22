from tkinter import *
from tkinter import messagebox, ttk
from tkcalendar import DateEntry
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import docx

import mysql.connector
import os
import random


mydb = mysql.connector.connect(host="localhost",user="root",passwd="my12345", database="tickets")
mycursor = mydb.cursor()

#mycursor.execute("Create table info8(T_Number int(20),P_Name varchar(200),P_Number int(20), Date varchar(200), From_ varchar(200), To_ varchar(200),Train varchar(200), N_of_Passenger int(20), T_Distance varchar(200), Amount varchar(200))")


document = Document()

if(mydb):
    print("Connection Successful")
else:
    print("Connection unsuccessful")


root = Tk()
user = Tk()
login = Tk()

bg_color = '#337442'
L_color = '#f72585'
F_color = 'f72585'

person = IntVar()
c_name = StringVar()
c_phone = StringVar()
ticket_no = StringVar()


login.title("Login")
login.geometry("300x200")
login.configure(background='#059DC0')
global entry1
global entry2
global entry3
Label(login, text="Username",font=('times new romon', 10, 'bold'), bg='#059DC0', fg='White').place(x=20, y=20)
Label(login, text="Password",font=('times new romon', 10, 'bold'), bg='#059DC0', fg='White').place(x=20, y=70)
entry1 = Entry(login, bd=5,font=('times new romon', 10, 'bold'))
entry1.place(x=140, y=20)
entry2 = Entry(login, bd=5, show= '*',font=('times new romon', 10, 'bold'))
entry2.place(x=140, y=70)



u1 = LabelFrame(user, bd=10, relief=GROOVE, text='Find Ticket', font=('times new romon', 15, 'bold'),fg='gold',bg=bg_color)
u1.place(x=0, y=80, relwidth=1)

Label(u1, text='Ticket Number', font=('times new romon', 18, 'bold'), bg=bg_color, fg='white').grid(row=0,column=0,padx=20,pady=5)
entry3 = Entry(u1, width=15,font='arial 15 bold', relief=SUNKEN, bd=7)
entry3.grid(row=0, column=1, padx=10,pady=5)

city = ['Banani railway station', 'Airport railway station', 'Cantonment railway station',
        'Gendaria railway station', 'Kamalapur railway station',
        'Shyampur Baraitala railway station', 'Tejgaon railway station']

train = ['Ekota Express','Parabat Express']

dict = {'Banani railway station- Airport railway station': 68.8,
        'Banani railway station- Cantonment railway station': 39.93,
        'Banani railway station-Gendaria railway station': 67.67,
        'Banani railway station-Kamalapur railway station': 25.07,
        'Banani railway station-Shyampur Baraitala railway station': 48.83,
        'Banani railway station-Tejgaon railway station': 99.65,
        'Airport railway station- Cantonment railway station': 51.86,
        'Airport railway station-Gendaria railway station': 54.55,
        'Airport railway station-Kamalapur railway station': 4.57,
        'Airport railway station-Shyampur Baraitala railway station': 6.18,
        'Airport railway station-Tejgaon railway station': 14.69,
        'Cantonment railway station-Gendaria railway station': 35.76,
        'Cantonment railway station-Kamalapur railway station': 22.95,
        'Cantonment railway station-Shyampur Baraitala railway station': 29.31,
        'Cantonment railway station-Tejgaon railway station': 58.24,
        'Gendaria railway station-Kamalapur railway station': 55.9,
        'Gendaria railway station-Shyampur Baraitala railway station': 72.12,
        'Gendaria railway station-Tejgaon railway station': 68.09,
        'Kamalapur railway station-Shyampur Baraitala railway station': 72.35,
        'Kamalapur railway station-Tejgaon railway station': 32.36,
        'Shyampur Baraitala railway station-Tejgaon railway station': 34.85
        }

def log():

    username = entry1.get()
    password = entry2.get()

    if username == "" and password == "":
        messagebox.showwarning("", "Blank Not Allowed")
    elif username == "Admin" and password == "G30":
        root.title('Admin Interface')
        root.geometry('1280x720')
        root.configure(background='#059DC0')
        login.destroy()
        user.destroy()

    elif username == "User" and password == "123":
        user.title('User Interface')
        user.geometry('600x650')
        user.configure(background='#059DC0')

        login.destroy()
        root.destroy()
    else:
        messagebox.showwarning("", "incorrect username and password")




def varify():
    global dis
    a = combo_s.get()
    b = combo_d.get()
    p = person.get()
    d=a+'-'+b
    e=b+'-'+a
    if c_name.get() != "" or c_phone.get() != "":
        if c_phone.get().isnumeric() is not True:
            messagebox.showerror('Error','Phone number should be integer')
            return
    else:
        messagebox.showerror("Error", "Passenger detail are must")
        return
    if a!=b:
        if d in dict:
            dis = dict[d]
        elif e in dict:
            dis = dict[e]
    else:
        messagebox.showwarning('Warning ','Please select right root')
        return
    messagebox.showinfo('Varified','Successfully Vairified')


def gticket():
   #try:
        welcome()
        p = person.get()
        textarea.insert(END, f"\n {55 * '*'}")
        textarea.insert(END, f"\n\n  From                 : {combo_s.get()}")
        textarea.insert(END, f"\n  To                     : {combo_d.get()}")
        textarea.insert(END, f"\n  Train                     : {combo_t.get()}")
        textarea.insert(END, f"\n  N. of Passenger : {p} person ")
        textarea.insert(END, f"\nTotal distance :\t\t{dis}")
        textarea.insert(END, f" km")
        textarea.insert(END, f"\n\n {35 * '='}")
        textarea.insert(END, f"\n  Amount :\t\t{2 * p * dis}")
        textarea.insert(END, f" Tk")
        textarea.insert(END, f"\n {35 * '='}")
        textarea.insert(END, f"\n\n {55 * '*'}")
        save_ticket()

        sqfrom = "Insert into info8(T_Number,P_Name,P_Number, Date, From_ , To_,Train, N_of_Passenger, T_Distance, Amount) values(%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)"
        info8 = [(ticket_no.get(), c_name.get(), c_phone.get(), cal.get(), combo_s.get(), combo_d.get(), combo_t.get(), person.get(), dis, 2 * person.get() * dis)]
        mycursor.executemany(sqfrom, info8)
        mydb.commit()

   #except Exception:
       # messagebox.showwarning('Warrning', 'Pleaes verify the details first')
       # clear()


def clear():
    c_name.set('')
    c_phone.set('')
    combo_s.set('select Source')
    combo_d.set('select destination')
    combo_t.set('select Train')
    person.set(0)
    welcome()

def exit():
    op = messagebox.askyesno("Exit", "Do you really want to exit?")
    if op > 0:
        root.destroy()

def save_ticket():
    op = messagebox.askyesno("Save ticket", "Do you want to download ticket ?")
    if op > 0:
        ticket_details = textarea.get('1.0', END)

        document.add_picture('D:\Consaltation\Ticket Generatorr\Files\logo.png')
        paragraph = document.add_paragraph(ticket_details)
        #paragraph.style = document.styles.add_style('Style Name', WD_STYLE_TYPE,PARAGRAPH)
        font = paragraph.style.font
        font.name ='Time New Roman'
        font.color.rgb = RGBColor(0, 215, 255)
        document.save(ticket_no.get()+'.docx')
        f1 = open("Ticket/" + str(ticket_no.get()) + ".txt", "w")
        f1.write(ticket_details)
       # f1.c
        f1.close()
        messagebox.showinfo("Saved", f"Ticket no, :{ticket_no.get()} Saved Successfully")

    else:
        return

def find_ticket():
    present="no"
    for i in os.listdir('Ticket/'):
        if i==str(entry3.get()) + ".txt":
            f1=open(f"Ticket/{i}","r")
            tarea.delete('1.0',END)
            for d in f1:
                tarea.insert(END,d)
            f1.close()
            present="yes"
    if present=="no":
        messagebox.showerror("Error",f"Ticket no: {entry3.get()} Not Found")

def welcome():
    x = random.randint(1000, 9999)
    ticket_no.set(str(x))
    textarea.delete(1.0, END)
    textarea.insert(END, "\t\t Have Safe Journey")
    textarea.insert(END, f"\n\n Ticket Number     : {ticket_no.get()}")
    textarea.insert(END, f"\n Passenger Name : {c_name.get()}")
    textarea.insert(END, f"\n Phone Number    : {c_phone.get()}")
    textarea.insert(END, f"\n Date                    : {cal.get()}")
    textarea.configure(font='arial 15 bold')

title = Label(root, pady=2, text="Intra Dhaka Railway Ticket Generator", bd=12, bg=bg_color, fg='white',font=('times new roman', 25, 'bold'), relief=GROOVE, justify=CENTER)
title.pack(fill=X)

title = Label(user, pady=2, text="Intra Dhaka Railway Ticket Check", bd=12, bg=bg_color, fg='white',font=('times new roman', 25, 'bold'), relief=GROOVE, justify=CENTER)
title.pack(fill=X)


F1 = LabelFrame(root, bd=10, relief=GROOVE, text='Passenger Details', font=('times new romon', 15, 'bold'),fg='gold',bg=bg_color)
F1.place(x=0, y=80, relwidth=1)

Label(F1, text='Passenger Name', font=('times new romon', 18, 'bold'), bg=bg_color, fg='white').grid(row=0,column=0,padx=20,pady=5)
Entry(F1, width=15, textvariable=c_name, font='arial 15 bold', relief=SUNKEN, bd=7).grid(row=0, column=1, padx=10,pady=5)

Label(F1, text='Phone No. ', font=('times new romon', 18, 'bold'), bg=bg_color, fg='white').grid(row=0, column=2,padx=20, pady=5)
Entry(F1, width=15, font='arial 15 bold', textvariable=c_phone, relief=SUNKEN, bd=7).grid(row=0, column=3, padx=10,pady=5)

Label(F1, text='Date', font=('times new romon', 18, 'bold'), bg=bg_color, fg='white').grid(row=0, column=4, padx=20,pady=5)

cal = DateEntry(F1, selectmode='day', font=('times new roman', 18), state='readonly')
cal.grid(row=0, column=5, padx=20, pady=5)

F2 = LabelFrame(root, text='Root Details', font=('times new romon', 18, 'bold'), fg='gold', bg=bg_color)
F2.place(x=20, y=180, width=630, height=600)

Label(F2, text='From', font=('times new roman', 18, 'bold'), bg=bg_color, fg='white').grid(row=0, column=0,padx=30, pady=20)
combo_s = ttk.Combobox(F2, font=('times new roman', 18), state='readonly', value=city)
combo_s.grid(row=0, column=1, pady=10)
combo_s.set('From Station')

Label(F2, text='To', font=('times new romon', 18, 'bold'), bg=bg_color, fg='white').grid(row=1, column=0, padx=10,pady=20)
combo_d = ttk.Combobox(F2, font=('times new roman', 18), state='readonly', value=city)
combo_d.grid(row=1, column=1, pady=10)
combo_d.set('To Station')

Label(F2, text='Train', font=('times new romon', 18, 'bold'), bg=bg_color, fg='white').grid(row=2, column=0, padx=10,pady=20)
combo_t = ttk.Combobox(F2, font=('times new roman', 18), state='readonly', value=train)
combo_t.grid(row=2, column=1, pady=10)
combo_t.set('train')

Label(F2, text='Number of ticket', font=('times new romon', 18, 'bold'), bg=bg_color, fg='white').grid(row=3,column=0,padx=30,pady=20)
Entry(F2, width=20, textvariable=person, font='arial 15 bold', relief=SUNKEN, bd=7).grid(row=3, column=1, padx=10,pady=20)



F3 = Frame(root, relief=GROOVE, bd=10)
F3.place(x=700, y=180, width=550, height=500)
Label(F3, text='Ticket', font='arial 15 bold', bd=7, relief=GROOVE).pack(fill=X)
scrol_y = Scrollbar(F3, orient=VERTICAL)
textarea = Text(F3, yscrollcommand=scrol_y)
scrol_y.pack(side=RIGHT, fill=Y)
scrol_y.config(command=textarea.yview)
textarea.pack()

welcome()



u2 = Frame(user, relief=GROOVE, bd=10)
u2.place(x=25, y=180, width=550, height=500)
Label(u2, text='Ticket', font='arial 15 bold', bd=7, relief=GROOVE).pack(fill=X)
scrol_y = Scrollbar(u2, orient=VERTICAL)
tarea = Text(u2, yscrollcommand=scrol_y.set)
scrol_y.pack(side=RIGHT, fill=Y)
scrol_y.config(command=tarea.yview)
tarea.pack()
welcome()



btn1 = Button(F2, text='Varify', font='arial 15 bold', command=varify, padx=5, pady=10, bg='#480CA8', width=15)
btn1.grid(row=4, column=0, padx=10, pady=30)
btn2 = Button(F2, text='Ticket', font='arial 15 bold', command=gticket, padx=5, pady=10, bg='#480CA8', width=15)
btn2.grid(row=4, column=1, padx=10, pady=30)
btn3 = Button(F2, text='Clear', font='arial 15 bold', padx=5, pady=10, command=clear, bg='#480CA8', width=15)
btn3.grid(row=5, column=0, padx=10, pady=30)
Button(F2, text='Exit', font='arial 15 bold', padx=5, pady=10, command=exit, bg='#480CA8', width=15).grid(row=5, column=1, padx=10, pady=30)
Button(login, text="Login", command=log, height=2, width=10, bd=6, font=('times new romon', 10, 'bold'), bg='#480CA8', fg='white').place(x=100, y=120)
Button(u1, width=10, command=find_ticket, text="Search", bg="white", fg="black", bd=7,font=("arial", 12, "bold")).grid(row=0, column=2, padx=10, pady=10)



user.mainloop()
login.mainloop()
root.mainloop()
